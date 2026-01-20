from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def add_title(doc, title_text):
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return title

def add_section_heading(doc, heading_text, level=1):
    return doc.add_heading(heading_text, level=level)

def add_paragraph(doc, text):
    return doc.add_paragraph(text)

def add_math_formula(doc, formula_text):
    p = doc.add_paragraph()
    run = p.add_run(formula_text)
    run.italic = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return p

def add_changelog_section(doc, version, changes):
    add_section_heading(doc, "Version History & Changelog")
    add_paragraph(doc, f"Version: {version}")
    add_paragraph(doc, "Updated to align with:")
    add_paragraph(doc, "  - DeepMamba v2.0 / v2.5 architecture")
    add_paragraph(doc, "  - Phase 2.5 lag-alignment system")
    add_paragraph(doc, "  - Topological Data Analysis (TDA) v2.0")
    add_paragraph(doc, "  - VolatilityEnergy & CurvatureProxy features")
    add_paragraph(doc, "")
    add_section_heading(doc, "Changes in This Version:", level=2)
    for change in changes:
        p = doc.add_paragraph(change, style='List Bullet')

def create_rule_c2_updated():
    """Rule C2: Persistent Homology Regime Shift (MAJOR UPDATE v2.5)"""
    doc = Document()
    add_title(doc, "Rule C2: Persistent Homology Regime Shift [v2.5 MAJOR UPDATE]")

    add_changelog_section(doc, "v2.5", [
        "MAJOR: β₁(t) now normalized using rolling z-score",
        "MAJOR: Added CurvatureProxy and VolatilityEnergy as gating modifiers",
        "MAJOR: Added minimum persistence threshold to filter noise",
        "MAJOR: Added lag-aware IV confidence weighting",
        "Complete rewrite of topological signature computation",
        "New parameter table with normalization coefficients"
    ])
    add_paragraph(doc, "")

    add_section_heading(doc, "Summary")
    add_paragraph(doc, "Uses topological regime signatures (persistent homology) to detect breakout or consolidation using Takens embedding and persistent diagram analysis. This v2.5 update introduces normalized β₁ scores, geometric curvature gating, and volatility energy modulation for superior regime detection.")
    add_paragraph(doc, "")
    add_paragraph(doc, "Strategy Type: Breakout / Topology-Based Regime Detection")
    add_paragraph(doc, "Timeframe: 5-minute, 15-minute (requires sufficient embedding dimension)")
    add_paragraph(doc, "Application: Regime shift detection, consolidation/breakout classification")

    add_section_heading(doc, "1. Mathematical Derivations (MAJOR REWRITE)")

    add_section_heading(doc, "1.1 Takens Embedding", level=2)
    add_paragraph(doc, "Construct delay-coordinate embedding of price time series:")
    add_math_formula(doc, "X(t) = [P(t), P(t-τ), P(t-2τ), ..., P(t-(d-1)τ)]")
    add_paragraph(doc, "")
    add_paragraph(doc, "Where:")
    add_paragraph(doc, "  - P(t) = price (or log-returns)")
    add_paragraph(doc, "  - τ = time delay (typically 1-3 bars)")
    add_paragraph(doc, "  - d = embedding dimension (typically 3-5)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Default: τ=1, d=3 for 5-minute data")

    add_section_heading(doc, "1.2 Persistent Homology Computation", level=2)
    add_paragraph(doc, "Apply Vietoris-Rips filtration to embedded point cloud:")
    add_math_formula(doc, "VR(ε) = {simplices with pairwise distance ≤ ε}")
    add_paragraph(doc, "")
    add_paragraph(doc, "Compute 1-dimensional Betti number (cycle count):")
    add_math_formula(doc, "β₁(ε) = number of independent cycles at scale ε")
    add_paragraph(doc, "")
    add_paragraph(doc, "Aggregate over filtration scales:")
    add_math_formula(doc, "β₁_raw(t) = max_ε β₁(ε, t)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Or use persistent entropy:")
    add_math_formula(doc, "β₁_entropy(t) = -Σᵢ pᵢ log(pᵢ)")
    add_paragraph(doc, "Where pᵢ = persistence_i / Σ persistence")

    add_section_heading(doc, "1.3 Normalized β₁ Score (NEW - CRITICAL)", level=2)
    add_paragraph(doc, "MAJOR UPDATE: β₁ must be normalized to be regime-invariant:")
    add_math_formula(doc, "β₁_norm(t) = (β₁_raw(t) - μ_β₁(w)) / σ_β₁(w)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Where:")
    add_paragraph(doc, "  - μ_β₁(w) = rolling mean of β₁_raw over window w")
    add_paragraph(doc, "  - σ_β₁(w) = rolling std of β₁_raw over window w")
    add_paragraph(doc, "  - Default: w = 100 bars")
    add_paragraph(doc, "")
    add_paragraph(doc, "Interpretation:")
    add_paragraph(doc, "  - β₁_norm > +2: High cycle count (consolidation, complex structure)")
    add_paragraph(doc, "  - β₁_norm < -1: Low cycle count (breakout, simple structure)")
    add_paragraph(doc, "  - |β₁_norm| < 0.5: Neutral regime")

    add_section_heading(doc, "1.4 Regime Shift Detection with Minimum Persistence (NEW)", level=2)
    add_paragraph(doc, "NEW: Filter noise using minimum persistence threshold:")
    add_paragraph(doc, "")
    add_paragraph(doc, "For each cycle (birth_i, death_i):")
    add_math_formula(doc, "persistence_i = death_i - birth_i")
    add_paragraph(doc, "")
    add_paragraph(doc, "Only count cycles where:")
    add_math_formula(doc, "persistence_i > θ_persist · max_persistence(t)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Default: θ_persist = 0.1 (ignore cycles with <10% max persistence)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Breakout detection (cycle destruction):")
    add_math_formula(doc, "Breakout(t) = (β₁_norm(t) < -1.0) AND (Δβ₁_norm(t) < -0.5)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Consolidation detection (cycle emergence):")
    add_math_formula(doc, "Consolidation(t) = (β₁_norm(t) > +2.0) AND (Δβ₁_norm(t) > +0.5)")

    add_section_heading(doc, "1.5 Curvature Proxy and Volatility Energy Gating (NEW)", level=2)
    add_paragraph(doc, "NEW: Modulate topological signals with geometric features:")
    add_paragraph(doc, "")
    add_paragraph(doc, "Curvature Proxy (from scientific_spec.md):")
    add_math_formula(doc, "κ(t) = |r(t+1) - 2·r(t) + r(t-1)| / Δt²")
    add_paragraph(doc, "Where r(t) = log-return at time t")
    add_paragraph(doc, "")
    add_paragraph(doc, "Volatility Energy:")
    add_math_formula(doc, "VolEnergy(t) = ATR(t) / SMA_50(ATR)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Gated topological score:")
    add_math_formula(doc, "β₁_gated(t) = β₁_norm(t) · (1 + α·κ(t)) · (1 + β·VolEnergy(t))")
    add_paragraph(doc, "")
    add_paragraph(doc, "Default: α = 0.2, β = 0.3")
    add_paragraph(doc, "")
    add_paragraph(doc, "This amplifies topological signals during high curvature / volatility regimes.")

    add_section_heading(doc, "1.6 Lag-Aware IV Confidence (NEW)", level=2)
    add_paragraph(doc, "NEW: Weight regime shifts by IV data freshness:")
    add_math_formula(doc, "IV_confidence(t) = exp(-λ · lag_minutes(t))")
    add_paragraph(doc, "")
    add_paragraph(doc, "Final regime shift score:")
    add_math_formula(doc, "RegimeScore(t) = β₁_gated(t) · IV_confidence(t)")

    add_section_heading(doc, "2. Logic Conditions (MAJOR REWRITE)")
    add_paragraph(doc, "For BREAKOUT entry:")
    add_paragraph(doc, "1. β₁_norm(t) < -1.0 (low cycle count)")
    add_paragraph(doc, "2. Δβ₁_norm(t) < -0.5 (sharp drop in cycles)")
    add_paragraph(doc, "3. Minimum persistence filter applied")
    add_paragraph(doc, "4. NEW: VolEnergy(t) > 1.0 (elevated volatility)")
    add_paragraph(doc, "5. NEW: IV_confidence(t) > 0.5")
    add_paragraph(doc, "6. NEW: |RegimeScore(t)| > 2.0 (strong signal)")
    add_paragraph(doc, "7. Consolidation score < 0.3 (from legacy rule)")
    add_paragraph(doc, "")
    add_paragraph(doc, "For CONSOLIDATION detection (avoid entries):")
    add_paragraph(doc, "1. β₁_norm(t) > +2.0")
    add_paragraph(doc, "2. Δβ₁_norm(t) > +0.5")
    add_paragraph(doc, "3. Regime filter: block new iron condor entries")

    add_section_heading(doc, "3. Entry and Exit Triggers")

    add_section_heading(doc, "3.1 Entry Trigger", level=2)
    add_paragraph(doc, "- Enter in direction of breakout when all 7 conditions met")
    add_paragraph(doc, "- Use as filter: only allow directional trades when breakout confirmed")

    add_section_heading(doc, "3.2 Exit Trigger", level=2)
    add_paragraph(doc, "- Exit when β₁_norm(t) rises (return to consolidation)")
    add_paragraph(doc, "- Exit when RegimeScore reverses sign")

    add_section_heading(doc, "4. Parameter Table (NEW)")
    table = doc.add_table(rows=11, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Default Value'
    hdr_cells[2].text = 'Description'

    params = [
        ('Embedding τ', '1', 'Time delay for Takens embedding'),
        ('Embedding d', '3', 'Embedding dimension'),
        ('β₁ Normalization w', '100', 'Rolling window for z-score'),
        ('Persistence Threshold θ', '0.1', 'Minimum persistence ratio'),
        ('Breakout β₁ Threshold', '-1.0', 'Z-score for breakout'),
        ('Consolidation β₁ Threshold', '+2.0', 'Z-score for consolidation'),
        ('Curvature Weight α', '0.2', 'Curvature gating coefficient'),
        ('VolEnergy Weight β', '0.3', 'Volatility energy coefficient'),
        ('IV Decay λ', '0.05', 'IV confidence decay rate'),
        ('RegimeScore Threshold', '2.0', 'Minimum gated score')
    ]

    for i, (param, value, desc) in enumerate(params, 1):
        row = table.rows[i].cells
        row[0].text = param
        row[1].text = value
        row[2].text = desc

    add_section_heading(doc, "5. Implementation Notes")
    add_paragraph(doc, "Python implementation (see intelligence/tda_signature.py):")
    add_paragraph(doc, "")
    add_paragraph(doc, "```python")
    add_paragraph(doc, "from ripser import ripser")
    add_paragraph(doc, "from sklearn.preprocessing import StandardScaler")
    add_paragraph(doc, "")
    add_paragraph(doc, "def compute_beta1_normalized(prices, tau=1, d=3, window=100):")
    add_paragraph(doc, "    # Takens embedding")
    add_paragraph(doc, "    X = takens_embed(prices, tau, d)")
    add_paragraph(doc, "    ")
    add_paragraph(doc, "    # Persistent homology")
    add_paragraph(doc, "    result = ripser(X, maxdim=1)")
    add_paragraph(doc, "    dgm1 = result['dgms'][1]  # 1-dimensional diagram")
    add_paragraph(doc, "    ")
    add_paragraph(doc, "    # Filter by persistence")
    add_paragraph(doc, "    persist = dgm1[:, 1] - dgm1[:, 0]")
    add_paragraph(doc, "    valid = persist > 0.1 * persist.max()")
    add_paragraph(doc, "    beta1_raw = valid.sum()")
    add_paragraph(doc, "    ")
    add_paragraph(doc, "    # Normalize")
    add_paragraph(doc, "    beta1_norm = (beta1_raw - rolling_mean) / rolling_std")
    add_paragraph(doc, "    return beta1_norm")
    add_paragraph(doc, "```")

    add_section_heading(doc, "6. Integration Notes")
    add_paragraph(doc, "- Persistent homology computed using ripser library")
    add_paragraph(doc, "- Takens embedding via intelligence/tda_signature.py")
    add_paragraph(doc, "- β₁ normalization critical for regime invariance")
    add_paragraph(doc, "- Curvature and VolEnergy from v2.1 feature schema")
    add_paragraph(doc, "- Compatible with CondorBrain regime classification head")

    add_section_heading(doc, "7. References")
    add_paragraph(doc, "- intelligence/tda_signature.py")
    add_paragraph(doc, "- scientific_spec.md (v2.1 topological features)")
    add_paragraph(doc, "- Ripser: https://ripser.scikit-tda.org/")

    return doc

def create_rule_d2_updated():
    """Rule D2: Volume Spike with Band Break (UPDATED v2.0)"""
    doc = Document()
    add_title(doc, "Rule D2: Volume Spike with Band Break [v2.0]")

    add_changelog_section(doc, "v2.0", [
        "Added volatility energy dampening to avoid false spikes",
        "Added lag-aware IV confidence to breakout confirmation",
        "Updated volume ratio normalization",
        "Added regime-aware volume threshold"
    ])
    add_paragraph(doc, "")

    add_section_heading(doc, "Summary")
    add_paragraph(doc, "Uses volume surge to confirm momentum breakout beyond dynamic bands. The rule validates institutional participation through regime-aware volume analysis.")
    add_paragraph(doc, "")
    add_paragraph(doc, "Strategy Type: Momentum / Volume Confirmation")
    add_paragraph(doc, "Timeframe: 1-minute, 5-minute")
    add_paragraph(doc, "Application: Momentum trades, breakout confirmation")

    add_section_heading(doc, "1. Mathematical Derivations")

    add_section_heading(doc, "1.1 Volume Spike Detection", level=2)
    add_paragraph(doc, "Volume ratio:")
    add_math_formula(doc, "VolumeRatio(t) = Volume(t) / SMA_20(Volume)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Spike condition:")
    add_math_formula(doc, "VolumeRatio(t) > 1.5")

    add_section_heading(doc, "1.2 Volatility Energy Dampening (NEW)", level=2)
    add_paragraph(doc, "NEW: Adjust volume threshold by volatility regime:")
    add_math_formula(doc, "VolEnergy(t) = ATR(t) / SMA_50(ATR)")
    add_math_formula(doc, "VolumeThreshold_dynamic(t) = 1.5 · (1 + γ · max(0, VolEnergy(t) - 1))")
    add_paragraph(doc, "")
    add_paragraph(doc, "Default: γ = 0.3")
    add_paragraph(doc, "")
    add_paragraph(doc, "This raises the volume bar during high volatility to filter noise.")

    add_section_heading(doc, "1.3 Band Break with IV Confidence (NEW)", level=2)
    add_paragraph(doc, "NEW: Weight breakout by IV data freshness:")
    add_math_formula(doc, "BandBreak(t) = 1 if Close(t) > Upper Band(t) OR Close(t) < Lower Band(t)")
    add_math_formula(doc, "IV_confidence(t) = exp(-λ · lag_minutes(t))")
    add_math_formula(doc, "Breakout_score(t) = BandBreak(t) · VolumeRatio(t) · IV_confidence(t)")

    add_section_heading(doc, "2. Logic Conditions (UPDATED)")
    add_paragraph(doc, "For entry:")
    add_paragraph(doc, "1. Close(t) breaks Upper/Lower Dynamic Band")
    add_paragraph(doc, "2. VolumeRatio(t) > VolumeThreshold_dynamic(t)")
    add_paragraph(doc, "3. NEW: IV_confidence(t) > 0.5")
    add_paragraph(doc, "4. NEW: Breakout_score(t) > 1.0")

    add_section_heading(doc, "3. Entry and Exit Triggers")
    add_section_heading(doc, "3.1 Entry Trigger", level=2)
    add_paragraph(doc, "- Enter in direction of breakout with volume confirmation")

    add_section_heading(doc, "3.2 Exit Trigger", level=2)
    add_paragraph(doc, "- Exit on volume drop (VolumeRatio < 1.0)")
    add_paragraph(doc, "- Exit on price reversal to middle band")

    add_section_heading(doc, "4. Parameter Table (UPDATED)")
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Default Value'
    hdr_cells[2].text = 'Description'

    params = [
        ('Base Volume Threshold', '1.5', 'Base volume ratio for spike'),
        ('VolEnergy Dampening γ', '0.3', 'Volatility dampening coefficient'),
        ('IV Confidence Threshold', '0.5', 'Minimum IV confidence'),
        ('Breakout Score Threshold', '1.0', 'Minimum composite score')
    ]

    for i, (param, value, desc) in enumerate(params, 1):
        row = table.rows[i].cells
        row[0].text = param
        row[1].text = value
        row[2].text = desc

    add_section_heading(doc, "5. Integration Notes")
    add_paragraph(doc, "- Volume threshold adapts to volatility regime")
    add_paragraph(doc, "- IV confidence prevents stale data entries")
    add_paragraph(doc, "- Simple but effective momentum confirmation")

    return doc

def create_rule_e2_updated():
    """Rule E2: Dynamic Band Width Expansion (UPDATED v2.0)"""
    doc = Document()
    add_title(doc, "Rule E2: Dynamic Band Width Expansion [v2.0]")

    add_changelog_section(doc, "v2.0", [
        "Expansion now percentile-based instead of fixed percentage",
        "Added MTF confirmation requirement",
        "Added lag-aware IV confidence",
        "Added volatility energy normalization"
    ])
    add_paragraph(doc, "")

    add_section_heading(doc, "Summary")
    add_paragraph(doc, "Uses dynamic Bollinger Band width expansion to detect volatility regime shifts and trend initiation.")
    add_paragraph(doc, "")
    add_paragraph(doc, "Strategy Type: Volatility-Based / Trend Detection")

    add_section_heading(doc, "1. Mathematical Derivations (UPDATED)")

    add_section_heading(doc, "1.1 Bandwidth Percentile (NEW)", level=2)
    add_paragraph(doc, "NEW: Use percentile ranking instead of absolute percentage:")
    add_math_formula(doc, "BW_percentile(t, w) = PercentileRank(Bandwidth(t), Bandwidth[t-w:t])")
    add_paragraph(doc, "")
    add_paragraph(doc, "Expansion detected when:")
    add_math_formula(doc, "BW_percentile(t, 100) > 75")

    add_section_heading(doc, "1.2 MTF Confirmation (NEW)", level=2)
    add_paragraph(doc, "NEW: Require multi-timeframe alignment:")
    add_math_formula(doc, "MTF_consensus(t) > 0.7 for bullish")

    add_section_heading(doc, "1.3 Volatility Energy Normalization (NEW)", level=2)
    add_paragraph(doc, "NEW: Normalize by volatility energy:")
    add_math_formula(doc, "Expansion_normalized(t) = BW_percentile(t) / (1 + VolEnergy(t))")

    add_section_heading(doc, "2. Logic Conditions (UPDATED)")
    add_paragraph(doc, "1. BW_percentile(t, 100) > 75")
    add_paragraph(doc, "2. NEW: MTF_consensus(t) > 0.7")
    add_paragraph(doc, "3. NEW: IV_confidence(t) > 0.5")
    add_paragraph(doc, "4. Price breaks upper/lower band")

    add_section_heading(doc, "3. Integration Notes")
    add_paragraph(doc, "- Percentile-based for regime invariance")
    add_paragraph(doc, "- MTF prevents false expansions")

    return doc

def create_rule_e3_updated():
    """Rule E3: Persistent Homology Chaos Detection (MAJOR UPDATE v2.5)"""
    doc = Document()
    add_title(doc, "Rule E3: Persistent Homology Chaos Detection [v2.5 MAJOR UPDATE]")

    add_changelog_section(doc, "v2.5", [
        "MAJOR: β₁ normalized using rolling z-score",
        "MAJOR: Added VolatilityEnergy and CurvatureProxy gating",
        "MAJOR: Fuzzy dampening instead of hard blocking",
        "Added minimum persistence threshold",
        "Complete rewrite of chaos detection logic"
    ])
    add_paragraph(doc, "")

    add_section_heading(doc, "Summary")
    add_paragraph(doc, "Uses persistent homology to detect chaotic, high-volatility regimes and apply fuzzy dampening to position sizing rather than hard blocking entries.")
    add_paragraph(doc, "")
    add_paragraph(doc, "Strategy Type: Volatility-Based / Risk Filter")

    add_section_heading(doc, "1. Mathematical Derivations (MAJOR REWRITE)")

    add_section_heading(doc, "1.1 Normalized β₁ for Chaos Detection", level=2)
    add_paragraph(doc, "MAJOR UPDATE: Same normalization as Rule C2:")
    add_math_formula(doc, "β₁_norm(t) = (β₁_raw(t) - μ_β₁(w)) / σ_β₁(w)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Chaos detection:")
    add_math_formula(doc, "Chaos(t) = (β₁_norm(t) > +2.5) AND (Δβ₁_norm(t) > +0.5)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Interpretation: Sharp rise in cycles = chaotic regime")

    add_section_heading(doc, "1.2 Fuzzy Dampening (NEW - CRITICAL)", level=2)
    add_paragraph(doc, "MAJOR UPDATE: Instead of blocking entries, apply fuzzy dampening:")
    add_math_formula(doc, "ChaosMembership(t) = sigmoid(β₁_norm(t) - 2.0)")
    add_math_formula(doc, "PositionSize_adjusted(t) = PositionSize_base · (1 - ChaosMembership(t))")
    add_paragraph(doc, "")
    add_paragraph(doc, "This smoothly reduces size as chaos increases, never hard-blocking.")

    add_section_heading(doc, "1.3 VolEnergy and Curvature Gating (NEW)", level=2)
    add_paragraph(doc, "NEW: Amplify chaos signal during high curvature:")
    add_math_formula(doc, "β₁_gated(t) = β₁_norm(t) · (1 + α·κ(t)) · (1 + β·VolEnergy(t))")
    add_paragraph(doc, "")
    add_paragraph(doc, "Default: α=0.2, β=0.3")

    add_section_heading(doc, "2. Logic Conditions (MAJOR REWRITE)")
    add_paragraph(doc, "For chaos dampening:")
    add_paragraph(doc, "1. β₁_gated(t) > +2.5 (high chaos)")
    add_paragraph(doc, "2. Apply ChaosMembership to reduce position size")
    add_paragraph(doc, "3. Resume normal operations when β₁_gated(t) < +1.5")

    add_section_heading(doc, "3. Integration Notes")
    add_paragraph(doc, "- Fuzzy dampening prevents hard blocks")
    add_paragraph(doc, "- Compatible with FIS position sizing")
    add_paragraph(doc, "- VolEnergy and curvature from v2.1 schema")

    return doc

def main():
    print("Generating MAJOR UPDATES for Rules C2, D2, E2, E3 (v2.0/v2.5)...")

    docs_dir = r'C:\SPYOptionTrader_test\docs\trading_rules\v2.0'
    os.makedirs(docs_dir, exist_ok=True)

    # Major Updates
    print("Creating Rule C2 (MAJOR UPDATE v2.5)...")
    doc_c2 = create_rule_c2_updated()
    doc_c2.save(os.path.join(docs_dir, 'Rule_C2_Persistent_Homology_v2.5.docx'))

    print("Creating Rule D2 (UPDATED v2.0)...")
    doc_d2 = create_rule_d2_updated()
    doc_d2.save(os.path.join(docs_dir, 'Rule_D2_Volume_Spike_v2.0.docx'))

    print("Creating Rule E2 (UPDATED v2.0)...")
    doc_e2 = create_rule_e2_updated()
    doc_e2.save(os.path.join(docs_dir, 'Rule_E2_Band_Width_Expansion_v2.0.docx'))

    print("Creating Rule E3 (MAJOR UPDATE v2.5)...")
    doc_e3 = create_rule_e3_updated()
    doc_e3.save(os.path.join(docs_dir, 'Rule_E3_Chaos_Detection_v2.5.docx'))

    print(f"\n=== PART A COMPLETE ===")
    print(f"\nAll 7 updated rules saved to: {docs_dir}")
    print("\nGenerated:")
    print("  MINOR UPDATES (v2.0):")
    print("    - Rule_A2_MACD_Crossover_v2.0.docx")
    print("    - Rule_B1_Fuzzy_Reversion_v2.0.docx")
    print("    - Rule_C1_Squeeze_Breakout_v2.0.docx")
    print("    - Rule_D2_Volume_Spike_v2.0.docx")
    print("    - Rule_E2_Band_Width_Expansion_v2.0.docx")
    print("  MAJOR UPDATES (v2.5):")
    print("    - Rule_C2_Persistent_Homology_v2.5.docx")
    print("    - Rule_E3_Chaos_Detection_v2.5.docx")
    print("\nPerfect rules (no changes needed):")
    print("    - Rule_A1, A3, B2, D1, E1, PSAR_Reversion")

if __name__ == "__main__":
    main()
