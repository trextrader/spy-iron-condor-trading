from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def add_title(doc, title_text):
    """Add a formatted title to the document"""
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return title

def add_section_heading(doc, heading_text, level=1):
    """Add a section heading"""
    heading = doc.add_heading(heading_text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph"""
    p = doc.add_paragraph(text)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    return p

def add_math_formula(doc, formula_text):
    """Add a mathematical formula"""
    p = doc.add_paragraph()
    run = p.add_run(formula_text)
    run.italic = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return p

def create_rule_a1():
    """Rule A1: Dynamic Bollinger Band Breakout with Volume Confirmation"""
    doc = Document()

    add_title(doc, "Rule A1: Dynamic Bollinger Band Breakout with Volume Confirmation")

    # Summary
    add_section_heading(doc, "Summary")
    add_paragraph(doc, "This rule identifies trend continuation opportunities using a volatility-adaptive Bollinger Band breakout combined with volume confirmation. The rule is designed to detect strong directional moves that emerge from low-volatility consolidation periods (band squeezes) and are confirmed by institutional participation (volume spikes).")
    add_paragraph(doc, "")
    add_paragraph(doc, "Strategy Type: Trend-Following")
    add_paragraph(doc, "Timeframe: 1-minute, 5-minute, 15-minute (Multi-Timeframe Consensus)")
    add_paragraph(doc, "Application: Long directional calls, call spreads, or iron condor exit signals")

    # Mathematical Derivations
    add_section_heading(doc, "1. Mathematical Derivations")

    add_section_heading(doc, "1.1 Dynamic Bollinger Bands", level=2)
    add_paragraph(doc, "Traditional Bollinger Bands use a fixed lookback period and standard deviation multiplier. Dynamic Bollinger Bands adapt to local volatility energy using geometric and topological features.")
    add_paragraph(doc, "")
    add_paragraph(doc, "Definition:")
    add_math_formula(doc, "Upper Band(t) = SMA_n(t) + k(t) · σ_n(t)")
    add_math_formula(doc, "Middle Band(t) = SMA_n(t)")
    add_math_formula(doc, "Lower Band(t) = SMA_n(t) − k(t) · σ_n(t)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Where:")
    add_paragraph(doc, "• SMA_n(t) = Simple Moving Average over n bars at time t")
    add_paragraph(doc, "• σ_n(t) = Standard deviation over n bars at time t")
    add_paragraph(doc, "• k(t) = Dynamic multiplier based on volatility regime")
    add_paragraph(doc, "")
    add_paragraph(doc, "The dynamic multiplier k(t) is computed as:")
    add_math_formula(doc, "k(t) = k_0 · (1 + α · VolatilityEnergy(t))")
    add_paragraph(doc, "")
    add_paragraph(doc, "Where k_0 is a base multiplier (typically 2.0) and α is a sensitivity parameter (typically 0.1-0.3).")

    add_section_heading(doc, "1.2 Bollinger Band Bandwidth (Squeeze Detection)", level=2)
    add_paragraph(doc, "The bandwidth measures the width of the Bollinger Bands relative to the middle band:")
    add_math_formula(doc, "Bandwidth(t) = (Upper Band(t) − Lower Band(t)) / Middle Band(t)")
    add_math_formula(doc, "Bandwidth(t) = (2 · k(t) · σ_n(t)) / SMA_n(t)")
    add_paragraph(doc, "")
    add_paragraph(doc, "A squeeze is detected when:")
    add_math_formula(doc, "Bandwidth(t) < Percentile_10(Bandwidth[t−w:t])")
    add_paragraph(doc, "")
    add_paragraph(doc, "Where w is a rolling window (typically 100-500 bars).")

    add_section_heading(doc, "1.3 Volume Confirmation", level=2)
    add_paragraph(doc, "Volume ratio confirms institutional participation:")
    add_math_formula(doc, "VolumeRatio(t) = Volume(t) / SMA_20(Volume[t−20:t])")
    add_paragraph(doc, "")
    add_paragraph(doc, "A volume spike is confirmed when:")
    add_math_formula(doc, "VolumeRatio(t) > 1.5")

    add_section_heading(doc, "1.4 ADX (Average Directional Index)", level=2)
    add_paragraph(doc, "The volatility-normalized ADX measures trend strength:")
    add_math_formula(doc, "ADX_normalized(t) = ADX(t) / (1 + β · VolatilityEnergy(t))")
    add_paragraph(doc, "")
    add_paragraph(doc, "Where β is a normalization coefficient (typically 0.1-0.2).")
    add_paragraph(doc, "")
    add_paragraph(doc, "Trend strength is confirmed when:")
    add_math_formula(doc, "ADX_normalized(t) > 25 AND ADX_normalized(t) > ADX_normalized(t−1)")

    # Logic Conditions
    add_section_heading(doc, "2. Logic Conditions")
    add_paragraph(doc, "The rule triggers when ALL of the following conditions are satisfied:")
    add_paragraph(doc, "")
    add_paragraph(doc, "1. Squeeze Condition: Bandwidth(t) < Percentile_10(Bandwidth[t−w:t])")
    add_paragraph(doc, "2. Breakout Condition: Close(t) > Upper Band(t)")
    add_paragraph(doc, "3. Volume Confirmation: VolumeRatio(t) > 1.5")
    add_paragraph(doc, "4. Trend Strength: ADX_normalized(t) > 25 AND ADX_normalized(t) > ADX_normalized(t−1)")

    # Entry and Exit Triggers
    add_section_heading(doc, "3. Entry and Exit Triggers")

    add_section_heading(doc, "3.1 Entry Trigger", level=2)
    add_paragraph(doc, "Enter LONG (call spread or directional call) at the close of the bar where:")
    add_paragraph(doc, "• All four logic conditions are satisfied")
    add_paragraph(doc, "• Optional: Multi-timeframe consensus confirms bullish alignment across 1m, 5m, 15m")

    add_section_heading(doc, "3.2 Exit Trigger", level=2)
    add_paragraph(doc, "Exit the position when ANY of the following occurs:")
    add_paragraph(doc, "• Price closes below Middle Band(t)")
    add_paragraph(doc, "• VolumeRatio(t) < 1.0 (volume drops below average)")
    add_paragraph(doc, "• ADX_normalized(t) begins to decline for 2 consecutive bars")
    add_paragraph(doc, "• Profit target or stop-loss is reached (as defined by risk management rules)")

    # Parameter Table
    add_section_heading(doc, "4. Parameter Table")
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Default Value'
    hdr_cells[2].text = 'Description'

    params = [
        ('n (BB Period)', '20', 'Lookback period for Bollinger Bands'),
        ('k_0 (Base Multiplier)', '2.0', 'Base standard deviation multiplier'),
        ('w (Squeeze Window)', '100', 'Rolling window for bandwidth percentile'),
        ('Volume Threshold', '1.5', 'Minimum volume ratio for confirmation'),
        ('ADX Threshold', '25', 'Minimum ADX for trend strength'),
        ('α (Volatility Sensitivity)', '0.2', 'Dynamic multiplier sensitivity')
    ]

    for i, (param, value, desc) in enumerate(params, 1):
        row = table.rows[i].cells
        row[0].text = param
        row[1].text = value
        row[2].text = desc

    # Integration Notes
    add_section_heading(doc, "5. Integration Notes")
    add_paragraph(doc, "• All thresholds are volatility-adaptive and regime-aware")
    add_paragraph(doc, "• The rule can be encoded as a feature vector for the CondorBrain Mamba-2 model")
    add_paragraph(doc, "• Compatible with fuzzy logic sizing engine (outputs confidence score 0-1)")
    add_paragraph(doc, "• Dynamic BB and ADX computations use geometric and topological features for regime invariance")
    add_paragraph(doc, "• Suitable for both rule-based execution and neural network training targets")

    # References
    add_section_heading(doc, "6. References")
    add_paragraph(doc, "• Dynamic Market Indicators - Math and Pythonic integration for CondorIntelligence")
    add_paragraph(doc, "• Scientific Specification: CondorBrain & Mamba Architecture (scientific_spec.md)")
    add_paragraph(doc, "• Institutional-Grade Trading Rules for SPY Options Model Integration")

    return doc

def create_rule_a2():
    """Rule A2: MACD Crossover with Dynamic Band Expansion"""
    doc = Document()

    add_title(doc, "Rule A2: MACD Crossover with Dynamic Band Expansion")

    add_section_heading(doc, "Summary")
    add_paragraph(doc, "This rule uses a volatility-adaptive MACD crossover in conjunction with dynamic Bollinger Band expansion to identify strong trend initiation. The rule captures the transition from consolidation (narrow bands) to trending (widening bands) with momentum confirmation from MACD.")
    add_paragraph(doc, "")
    add_paragraph(doc, "Strategy Type: Trend-Following / Momentum")
    add_paragraph(doc, "Timeframe: 1-minute, 5-minute, 15-minute")
    add_paragraph(doc, "Application: Long/short directional trades, call/put spreads")

    add_section_heading(doc, "1. Mathematical Derivations")

    add_section_heading(doc, "1.1 Volatility-Adaptive MACD", level=2)
    add_paragraph(doc, "Standard MACD is computed as:")
    add_math_formula(doc, "MACD(t) = EMA_12(t) − EMA_26(t)")
    add_math_formula(doc, "Signal(t) = EMA_9(MACD(t))")
    add_math_formula(doc, "Histogram(t) = MACD(t) − Signal(t)")
    add_paragraph(doc, "")
    add_paragraph(doc, "The volatility-normalized MACD divides by realized volatility:")
    add_math_formula(doc, "MACD_normalized(t) = MACD(t) / σ_n(t)")
    add_math_formula(doc, "Signal_normalized(t) = Signal(t) / σ_n(t)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Where σ_n(t) is the n-period standard deviation of returns.")

    add_section_heading(doc, "1.2 Dynamic Bollinger Band Expansion", level=2)
    add_paragraph(doc, "Band expansion is measured by the percentage change in bandwidth:")
    add_math_formula(doc, "Expansion(t, y) = (Bandwidth(t) − Bandwidth(t−y)) / Bandwidth(t−y) × 100%")
    add_paragraph(doc, "")
    add_paragraph(doc, "An expansion event is detected when:")
    add_math_formula(doc, "Expansion(t, y) > X%")
    add_paragraph(doc, "")
    add_paragraph(doc, "Default: X = 10%, y = 5 bars")

    add_section_heading(doc, "1.3 MACD Crossover Condition", level=2)
    add_paragraph(doc, "Bullish crossover:")
    add_math_formula(doc, "MACD_normalized(t) > Signal_normalized(t) AND MACD_normalized(t−1) ≤ Signal_normalized(t−1)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Bearish crossover:")
    add_math_formula(doc, "MACD_normalized(t) < Signal_normalized(t) AND MACD_normalized(t−1) ≥ Signal_normalized(t−1)")

    add_section_heading(doc, "2. Logic Conditions")
    add_paragraph(doc, "For LONG entry:")
    add_paragraph(doc, "1. Bollinger Bands widening: Expansion(t, 5) > 10%")
    add_paragraph(doc, "2. MACD bullish crossover: MACD_normalized(t) crosses above Signal_normalized(t)")
    add_paragraph(doc, "3. ADX_normalized(t) > 20 (minimum trend strength)")
    add_paragraph(doc, "4. Price breaks above Upper Dynamic Band")
    add_paragraph(doc, "")
    add_paragraph(doc, "For SHORT entry:")
    add_paragraph(doc, "1. Bollinger Bands widening: Expansion(t, 5) > 10%")
    add_paragraph(doc, "2. MACD bearish crossover: MACD_normalized(t) crosses below Signal_normalized(t)")
    add_paragraph(doc, "3. ADX_normalized(t) > 20")
    add_paragraph(doc, "4. Price breaks below Lower Dynamic Band")

    add_section_heading(doc, "3. Entry and Exit Triggers")

    add_section_heading(doc, "3.1 Entry Trigger", level=2)
    add_paragraph(doc, "• LONG: Enter at bar close when all four logic conditions are met")
    add_paragraph(doc, "• SHORT: Enter at bar close when all four logic conditions are met")

    add_section_heading(doc, "3.2 Exit Trigger", level=2)
    add_paragraph(doc, "Exit when:")
    add_paragraph(doc, "• MACD line crosses back (reversal crossover)")
    add_paragraph(doc, "• Price closes below/above Middle Dynamic Band (for long/short)")
    add_paragraph(doc, "• Bollinger Band width begins to contract (Expansion < 0%)")

    add_section_heading(doc, "4. Parameter Table")
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Default Value'
    hdr_cells[2].text = 'Description'

    params = [
        ('MACD Fast Period', '12', 'Fast EMA period'),
        ('MACD Slow Period', '26', 'Slow EMA period'),
        ('MACD Signal Period', '9', 'Signal line EMA period'),
        ('Expansion Threshold X', '10%', 'Minimum bandwidth expansion'),
        ('Expansion Lookback y', '5', 'Bars to measure expansion over')
    ]

    for i, (param, value, desc) in enumerate(params, 1):
        row = table.rows[i].cells
        row[0].text = param
        row[1].text = value
        row[2].text = desc

    add_section_heading(doc, "5. Integration Notes")
    add_paragraph(doc, "• MACD and BB expansion are computed using volatility-adaptive formulas")
    add_paragraph(doc, "• Suitable for direct execution and as model training features")
    add_paragraph(doc, "• Compatible with fuzzy logic confidence scoring")
    add_paragraph(doc, "• Crossover signals can be pre-computed for faster inference")

    add_section_heading(doc, "6. References")
    add_paragraph(doc, "• Dynamic Market Indicators - Math and Pythonic integration")
    add_paragraph(doc, "• Scientific Specification: CondorBrain & Mamba Architecture")

    return doc

def create_rule_a3():
    """Rule A3: ADX-RSI Trend Confirmation"""
    doc = Document()

    add_title(doc, "Rule A3: ADX-RSI Trend Confirmation")

    add_section_heading(doc, "Summary")
    add_paragraph(doc, "This rule confirms trend strength using the Average Directional Index (ADX) and times entries using the Relative Strength Index (RSI). It captures strong momentum moves in the direction of the prevailing trend, entering when RSI reaches extreme levels that confirm momentum rather than exhaustion.")
    add_paragraph(doc, "")
    add_paragraph(doc, "Strategy Type: Trend-Following / Momentum")
    add_paragraph(doc, "Timeframe: 1-minute, 5-minute, 15-minute")
    add_paragraph(doc, "Application: Directional calls/puts, momentum spreads")

    add_section_heading(doc, "1. Mathematical Derivations")

    add_section_heading(doc, "1.1 Volatility-Normalized ADX", level=2)
    add_paragraph(doc, "The ADX measures trend strength on a 0-100 scale. Standard ADX computation:")
    add_math_formula(doc, "True Range (TR) = max(H(t)−L(t), |H(t)−C(t−1)|, |L(t)−C(t−1)|)")
    add_math_formula(doc, "+DM = max(H(t)−H(t−1), 0) if H(t)−H(t−1) > L(t−1)−L(t), else 0")
    add_math_formula(doc, "−DM = max(L(t−1)−L(t), 0) if L(t−1)−L(t) > H(t)−H(t−1), else 0")
    add_math_formula(doc, "+DI = 100 × Smooth_14(+DM) / Smooth_14(TR)")
    add_math_formula(doc, "−DI = 100 × Smooth_14(−DM) / Smooth_14(TR)")
    add_math_formula(doc, "DX = 100 × |+DI − −DI| / |+DI + −DI|")
    add_math_formula(doc, "ADX = Smooth_14(DX)")
    add_paragraph(doc, "")
    add_paragraph(doc, "Volatility normalization:")
    add_math_formula(doc, "ADX_normalized(t) = ADX(t) / (1 + β · VolatilityEnergy(t))")

    add_section_heading(doc, "1.2 Dynamic RSI", level=2)
    add_paragraph(doc, "Dynamic RSI is weighted by volatility energy and geometric curvature:")
    add_math_formula(doc, "RS(t) = Average_Gain_14(t) / Average_Loss_14(t)")
    add_math_formula(doc, "RSI(t) = 100 − (100 / (1 + RS(t)))")
    add_paragraph(doc, "")
    add_paragraph(doc, "The dynamic version adjusts thresholds based on regime:")
    add_math_formula(doc, "RSI_dynamic(t) = RSI(t) × (1 + γ · CurvatureProxy(t))")

    add_section_heading(doc, "1.3 Trend Direction and RSI Alignment", level=2)
    add_paragraph(doc, "For momentum LONG in an uptrend:")
    add_math_formula(doc, "+DI(t) > −DI(t) AND RSI_dynamic(t) > 70")
    add_paragraph(doc, "")
    add_paragraph(doc, "For momentum SHORT in a downtrend:")
    add_math_formula(doc, "−DI(t) > +DI(t) AND RSI_dynamic(t) < 30")

    add_section_heading(doc, "2. Logic Conditions")
    add_paragraph(doc, "For LONG entry (momentum uptrend):")
    add_paragraph(doc, "1. ADX_normalized(t) > 25")
    add_paragraph(doc, "2. ADX_normalized(t) > ADX_normalized(t−1) (rising ADX)")
    add_paragraph(doc, "3. +DI(t) > −DI(t)")
    add_paragraph(doc, "4. RSI_dynamic(t) > 70")
    add_paragraph(doc, "")
    add_paragraph(doc, "For SHORT entry (momentum downtrend):")
    add_paragraph(doc, "1. ADX_normalized(t) > 25")
    add_paragraph(doc, "2. ADX_normalized(t) > ADX_normalized(t−1)")
    add_paragraph(doc, "3. −DI(t) > +DI(t)")
    add_paragraph(doc, "4. RSI_dynamic(t) < 30")

    add_section_heading(doc, "3. Entry and Exit Triggers")

    add_section_heading(doc, "3.1 Entry Trigger", level=2)
    add_paragraph(doc, "• Enter at bar close when all four logic conditions are satisfied")
    add_paragraph(doc, "• Optional: Confirm with multi-timeframe consensus")

    add_section_heading(doc, "3.2 Exit Trigger", level=2)
    add_paragraph(doc, "Exit when:")
    add_paragraph(doc, "• ADX peaks and declines for 2 consecutive bars")
    add_paragraph(doc, "• RSI shows divergence (price makes new high/low but RSI doesn't)")
    add_paragraph(doc, "• RSI exits extreme zone (crosses back above 30 for short, below 70 for long)")
    add_paragraph(doc, "• +DI and −DI crossover (trend reversal)")

    add_section_heading(doc, "4. Parameter Table")
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Default Value'
    hdr_cells[2].text = 'Description'

    params = [
        ('ADX Period', '14', 'Smoothing period for ADX calculation'),
        ('ADX Threshold', '25', 'Minimum ADX for trend confirmation'),
        ('RSI Period', '14', 'Lookback period for RSI'),
        ('RSI Overbought', '70', 'Upper threshold for momentum long')
    ]

    for i, (param, value, desc) in enumerate(params, 1):
        row = table.rows[i].cells
        row[0].text = param
        row[1].text = value
        row[2].text = desc

    add_section_heading(doc, "5. Integration Notes")
    add_paragraph(doc, "• Dynamic RSI and ADX are volatility-normalized for regime invariance")
    add_paragraph(doc, "• Particularly effective in strong trend regimes")
    add_paragraph(doc, "• Can be combined with other momentum filters")
    add_paragraph(doc, "• Compatible with CondorBrain feature pipeline")

    return doc

# Create all rule documents
def main():
    print("Generating trading rule documents...")

    # Create docs directory if it doesn't exist
    docs_dir = r'C:\SPYOptionTrader_test\docs\trading_rules'
    os.makedirs(docs_dir, exist_ok=True)

    # Rule A1
    print("Creating Rule A1...")
    doc_a1 = create_rule_a1()
    doc_a1.save(os.path.join(docs_dir, 'Rule_A1_Dynamic_Bollinger_Band_Breakout.docx'))

    # Rule A2
    print("Creating Rule A2...")
    doc_a2 = create_rule_a2()
    doc_a2.save(os.path.join(docs_dir, 'Rule_A2_MACD_Crossover_with_Band_Expansion.docx'))

    # Rule A3
    print("Creating Rule A3...")
    doc_a3 = create_rule_a3()
    doc_a3.save(os.path.join(docs_dir, 'Rule_A3_ADX_RSI_Trend_Confirmation.docx'))

    print(f"\nCompleted! Documents saved to: {docs_dir}")
    print("Generated 3 rule documents (A1, A2, A3)")
    print("\nNote: Additional rules (B1-E3) will be generated in subsequent runs.")

if __name__ == "__main__":
    main()
