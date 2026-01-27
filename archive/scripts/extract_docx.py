from docx import Document
import sys

# Extract institutional trading rules
doc1 = Document(r'C:\SPYOptionTrader_test\docs\institutional-Grade Initial set of Trading Rules for SPY Options Model Integration.docx')
with open(r'C:\SPYOptionTrader_test\docs\institutional_rules_extracted.txt', 'w', encoding='utf-8') as f:
    for p in doc1.paragraphs:
        f.write(p.text + '\n')

# Extract template
doc2 = Document(r'C:\SPYOptionTrader_test\docs\The Spread -Average of H-L trading rule.docx')
with open(r'C:\SPYOptionTrader_test\docs\template_extracted.txt', 'w', encoding='utf-8') as f:
    for p in doc2.paragraphs:
        f.write(p.text + '\n')

print("Extraction complete!")
