from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def add_title(doc, text):
    title = doc.add_heading(text, level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 51, 102)
    return title

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph(text)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    # Light gray background
    shading_elm = p._element.get_or_add_pPr()
    return p

def add_warning_box(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("⚠️ " + text)
    run.bold = True
    run.font.color.rgb = RGBColor(153, 0, 0)
    p.paragraph_format.left_indent = Inches(0.25)
    return p

def create_antigrav_instructions():
    doc = Document()

    # Title
    add_title(doc, "Directive for AntiGrav (Claude Opus 4.5)")
    add_paragraph(doc, "Rule Implementation & Execution Decision Logic")
    add_paragraph(doc, "Version: 2.5 | Date: 2026-01-19")
    add_paragraph(doc, "")

    # Introduction
    add_heading(doc, "Introduction")
    add_paragraph(doc, "Hey AntiGrav,")
    add_paragraph(doc, "")
    add_paragraph(doc, "I've completed a comprehensive rule documentation and primitive extraction project using Claude Sonnet 4.5. Your task is to consume these artifacts and build the execution decision logic for the CondorBrain model training pipeline.")
    add_paragraph(doc, "")

    # Context
    add_heading(doc, "What Was Done (Context)")
    add_paragraph(doc, "13 institutional trading rules were:")
    add_bullet(doc, "Updated to align with DeepMamba v2.5, Phase 2.5 lag-alignment, and the 11-factor FIS system")
    add_bullet(doc, "Decomposed into 52 reusable primitives (instead of thousands of brittle if/then statements)")
    add_bullet(doc, "Encoded in a declarative YAML Rule DSL for systematic ingestion")
    add_paragraph(doc, "")

    # File Structure
    add_heading(doc, "Where to Look (File Locations)")
    add_paragraph(doc, "All documentation is located in the following structure:")
    add_paragraph(doc, "")

    add_code_block(doc, """📁 docs/
├── trading_rules/
│   ├── v2.0/                                  [7 updated .docx files]
│   │   ├── Rule_A2_MACD_Crossover_v2.0.docx
│   │   ├── Rule_B1_Fuzzy_Reversion_v2.0.docx
│   │   ├── Rule_C1_Squeeze_Breakout_v2.0.docx
│   │   ├── Rule_C2_Persistent_Homology_v2.5.docx    [MAJOR UPDATE]
│   │   ├── Rule_D2_Volume_Spike_v2.0.docx
│   │   ├── Rule_E2_Band_Width_Expansion_v2.0.docx
│   │   └── Rule_E3_Chaos_Detection_v2.5.docx        [MAJOR UPDATE]
│   │
│   └── [13 original .docx files]             [6 perfect, 7 superseded]
│       Rule_A1, Rule_A3, Rule_B2, Rule_D1,
│       Rule_E1, Example_Rule_PSAR_Reversion
│
├── Canonical_Rule_Primitive_Library.md        [52 primitives]
├── Rule_DSL_Specification.yaml                [Schema + 4 examples]
├── Complete_Ruleset_DSL.yaml                  [All 13 rules encoded]
└── IMPLEMENTATION_SUMMARY.md                  [Overview + metrics]""")

    add_paragraph(doc, "")

    # Reading Order
    add_heading(doc, "What to Do (Your Tasks)")
    add_heading(doc, "1. START HERE: Read These in Order", level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Order'
    hdr_cells[1].text = 'Document'
    hdr_cells[2].text = 'Purpose'

    reading_order = [
        ('1', 'IMPLEMENTATION_SUMMARY.md', 'Get the big picture (5 min read)'),
        ('2', 'Canonical_Rule_Primitive_Library.md', 'Understand the 52 building blocks (15 min)'),
        ('3', 'Rule_DSL_Specification.yaml', 'Learn DSL syntax and see 4 examples (10 min)'),
        ('4', 'Complete_Ruleset_DSL.yaml', 'See all 13 rules encoded (reference)')
    ]

    for i, (order, doc_name, purpose) in enumerate(reading_order, 1):
        row = table.rows[i].cells
        row[0].text = order
        row[1].text = doc_name
        row[2].text = purpose

    add_paragraph(doc, "")

    # Primary Objective
    add_heading(doc, "2. PRIMARY OBJECTIVE: Build Execution Decision Logic", level=2)
    add_paragraph(doc, "Your goal is to make the model learn these rules correctly by implementing:")
    add_paragraph(doc, "")

    # Task A
    add_heading(doc, "A. Primitive Computation Module", level=3)
    add_paragraph(doc, "Location: intelligence/primitives/", italic=True)
    add_bullet(doc, "Implement all 52 primitives as Python functions")
    add_bullet(doc, "Each primitive must match the mathematical formulas in Canonical_Rule_Primitive_Library.md")
    add_bullet(doc, "Ensure all primitives output the exact schema defined (e.g., {upper, middle, lower, bandwidth} for P001)")
    add_paragraph(doc, "")

    # Task B
    add_heading(doc, "B. DSL Parser", level=3)
    add_paragraph(doc, "Location: intelligence/rule_engine/dsl_parser.py", italic=True)
    add_bullet(doc, "Parse Complete_Ruleset_DSL.yaml")
    add_bullet(doc, "Build primitive dependency graph")
    add_bullet(doc, "Validate logical expressions (AND, OR, SEQ, THRESHOLD, etc.)")
    add_bullet(doc, "Topologically sort primitives for efficient computation")
    add_paragraph(doc, "")

    # Task C
    add_heading(doc, "C. Rule Execution Engine", level=3)
    add_paragraph(doc, "Location: intelligence/rule_engine/executor.py", italic=True)
    add_paragraph(doc, "Implement the 6-phase execution flow from Rule_DSL_Specification.yaml:")
    add_paragraph(doc, "")

    add_code_block(doc, """1. Precompute Primitives (parallelized)
2. Evaluate Signal Logic
3. Apply Gate Stack (short-circuit on BLOCK)
4. Compute Position Sizing
5. Risk Check
6. Execute/Log""")

    add_paragraph(doc, "")

    # Task D
    add_heading(doc, "D. Feature Pipeline Integration", level=3)
    add_bullet(doc, "Map primitive outputs to v2.1 feature columns")
    add_bullet(doc, "Ensure primitives are precomputed during data load (not at inference time)")
    add_bullet(doc, "Cache results for O(1) lookup")
    add_paragraph(doc, "")

    # Task E
    add_heading(doc, "E. Model Training Integration", level=3)
    add_bullet(doc, "Export primitive outputs as training features")
    add_bullet(doc, "Train auxiliary heads for gate outputs (ALLOW/BLOCK decisions)")
    add_bullet(doc, "Implement fuzzy constraint layers (11-factor memberships → neural output constraints)")
    add_paragraph(doc, "")

    # Critical Notes
    add_heading(doc, "Critical Notes")

    add_warning_box(doc, "MAJOR UPDATES (Rules C2 & E3):")
    add_bullet(doc, "These use normalized β₁ (persistent homology z-scores) - NOT raw β₁")
    add_bullet(doc, "Rule E3 uses fuzzy dampening (smooth 0-100% size reduction via sigmoid) - NOT hard blocking")
    add_bullet(doc, "Both require CurvatureProxy and VolatilityEnergy gating")
    add_paragraph(doc, "")

    add_warning_box(doc, "11-Factor FIS (Rule B1):")
    add_bullet(doc, "Fuzzy consensus weights: MTF=0.25, IVR=0.15, VIX=0.10, RSI=0.15, Stoch=0.05, ADX=0.05, SMA=0.05, PSAR=0.10, BB=0.05, BBsqueeze=0.03, Vol=0.02")
    add_bullet(doc, "Membership functions are in Canonical_Rule_Primitive_Library.md (M001-M011)")
    add_paragraph(doc, "")

    add_warning_box(doc, "Execution Friction Gate (Rule E1):")
    add_bullet(doc, "This is THE institutional execution filter - prevents bad fills")
    add_bullet(doc, "Formula: F(t) = Spread(t) / AvgRange(n, t), block if F >= θ(t)")
    add_bullet(doc, "Gap risk override: forces exit when G(t) >= 0.7 regardless of friction")
    add_paragraph(doc, "")

    # Validation Checklist
    add_heading(doc, "Validation Checklist")
    add_paragraph(doc, "Before marking complete, ensure:")
    add_paragraph(doc, "")

    checklist_items = [
        "All 52 primitives implemented and unit-tested",
        "DSL parser successfully loads Complete_Ruleset_DSL.yaml",
        "Execution engine runs all 13 rules without errors",
        "Primitive outputs match expected feature schema",
        "Gate stack short-circuits correctly (stops at first BLOCK)",
        "Fuzzy dampening works (E3 chaos → smooth size reduction)",
        "Backtest comparison: rule signals match historical expectations",
        "Model training: auxiliary targets (gates, regime) converge"
    ]

    for item in checklist_items:
        p = doc.add_paragraph(item, style='List Bullet')
        # Add checkbox
        p.insert_paragraph_before('☐ ' + item)
        p._element.getparent().remove(p._element)

    add_paragraph(doc, "")

    # Questions to Answer
    add_heading(doc, "Questions to Answer Back")
    add_paragraph(doc, "When you're done, report:")
    add_paragraph(doc, "")

    add_paragraph(doc, "1. Implementation status: Which primitives/modules are complete?", bold=True)
    add_paragraph(doc, "2. Blockers: Any missing data, unclear formulas, or architectural conflicts?", bold=True)
    add_paragraph(doc, "3. Validation results: Did rules backtest as expected? Any discrepancies?", bold=True)
    add_paragraph(doc, "4. Model integration: Are auxiliary heads training correctly? Loss curves?", bold=True)
    add_paragraph(doc, "")

    # Why This Matters
    add_heading(doc, "Why This Matters")
    add_paragraph(doc, "This primitive library eliminates the \"thousands of brittle if/then statements\" problem. Instead of hardcoding rules, you're building a compositional rule engine that:")
    add_paragraph(doc, "")

    add_bullet(doc, "Can add new rules without touching code (just edit YAML)")
    add_bullet(doc, "Exports clean features for model training")
    add_bullet(doc, "Provides interpretable auxiliary targets (gates, regime classifiers)")
    add_bullet(doc, "Enforces institutional risk discipline via fuzzy constraints")
    add_paragraph(doc, "")

    add_paragraph(doc, "The model will learn from these rules (via feature engineering and auxiliary tasks) while the rule engine provides safety rails (gates, constraints).", italic=True)
    add_paragraph(doc, "")

    # Primitive Summary Table
    add_heading(doc, "Quick Reference: Primitive Categories")

    prim_table = doc.add_table(rows=6, cols=3)
    prim_table.style = 'Light Grid Accent 1'

    hdr = prim_table.rows[0].cells
    hdr[0].text = 'Category'
    hdr[1].text = 'Count'
    hdr[2].text = 'ID Range'

    primitives = [
        ('Core Computational', '12', 'P001-P012'),
        ('Signal Generation', '15', 'S001-S015'),
        ('Gate Primitives', '10', 'G001-G010'),
        ('Membership Functions', '11', 'M001-M011'),
        ('Composite Primitives', '4', 'C001-C004')
    ]

    for i, (cat, count, id_range) in enumerate(primitives, 1):
        row = prim_table.rows[i].cells
        row[0].text = cat
        row[1].text = count
        row[2].text = id_range

    add_paragraph(doc, "")
    add_paragraph(doc, "Total Primitives: 52", bold=True)
    add_paragraph(doc, "")

    # Rule Status Summary
    add_heading(doc, "Quick Reference: Rule Update Status")

    status_table = doc.add_table(rows=4, cols=3)
    status_table.style = 'Light Grid Accent 1'

    hdr = status_table.rows[0].cells
    hdr[0].text = 'Status'
    hdr[1].text = 'Count'
    hdr[2].text = 'Rules'

    rule_status = [
        ('Perfect (no changes)', '6', 'A1, A3, B2, D1, E1, PSAR'),
        ('Minor Updates (v2.0)', '5', 'A2, B1, C1, D2, E2'),
        ('Major Updates (v2.5)', '2', 'C2, E3')
    ]

    for i, (status, count, rules) in enumerate(rule_status, 1):
        row = status_table.rows[i].cells
        row[0].text = status
        row[1].text = count
        row[2].text = rules

    add_paragraph(doc, "")

    # Closing
    add_heading(doc, "Next Steps")
    add_paragraph(doc, "Ready to start? Begin with IMPLEMENTATION_SUMMARY.md and let me know if you need clarification on any primitive definitions or DSL syntax.", bold=True)
    add_paragraph(doc, "")
    add_paragraph(doc, "All files are located in: C:\\SPYOptionTrader_test\\docs\\", italic=True)
    add_paragraph(doc, "")
    add_paragraph(doc, "Good luck, and reach out if you hit any blockers!")
    add_paragraph(doc, "")
    add_paragraph(doc, "")
    add_paragraph(doc, "---")
    add_paragraph(doc, "Generated by Claude Sonnet 4.5 | 2026-01-19", italic=True)

    return doc

def main():
    print("Generating AntiGrav Instructions Document...")

    doc = create_antigrav_instructions()

    output_path = r'C:\SPYOptionTrader_test\docs\trading_rules\AntiGrav_Implementation_Instructions.docx'
    doc.save(output_path)

    print(f"\nDocument created successfully!")
    print(f"Location: {output_path}")
    print("\nDocument includes:")
    print("  - Complete file structure and reading order")
    print("  - 5 implementation tasks (A-E)")
    print("  - Critical notes on major updates")
    print("  - Validation checklist")
    print("  - Quick reference tables")
    print("  - Questions to answer back")

if __name__ == "__main__":
    main()
